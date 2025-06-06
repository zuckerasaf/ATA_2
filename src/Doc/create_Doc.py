"""
Document generation utilities for test results.

This module provides functionality to generate Word (.docx) documents from JSON test data, including step tables and images.

Functions
---------
create_doc_from_json(json_path, pictures=True, Type="ATP", Regular_doc_path=True)
    Generate a Word document from a JSON test result file, including step tables and images.
"""

import json
import sys
import os
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
from src.utils.config import Config
import subprocess

# Add project root to Python path
project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), "../.."))
sys.path.insert(0, project_root)

def load_doc_config():
    """
    Load the document configuration from Doc_config.json.
    
    This function reads the configuration file that contains system information,
    environment settings, and document formatting preferences.
    
    Returns:
        dict: The configuration data containing:
            - system_info: Software and hardware details
            - environment: Test environment settings
            - document_settings: Formatting and styling preferences
    """
    try:
        config_path = os.path.join(os.path.dirname(__file__), "Doc_config.json")
        with open(config_path, 'r') as f:
            return json.load(f)
    except Exception as e:
        print(f"Error loading document config: {e}")
        return {}

def add_system_info_section(doc, config):
    """
    Add system information section to the document.
    
    This function creates a detailed section about the system configuration,
    including both software and hardware information. It formats the information
    in a hierarchical structure with appropriate headings.
    
    Args:
        doc: The document object to add content to
        config: The configuration dictionary containing system information
    
    The section includes:
        - Software Information:
            * Multiple software components (SW_1, SW_2, etc.)
            * Version and comments for each
        - Hardware Information:
            * Multiple PC configurations (PC_1, PC_2, etc.)
            * Comments and specifications for each
    """
    doc.add_heading("System Information", level=1)
    
    # Software Information
    doc.add_heading("Software", level=2)
    software = config.get('system_info', {}).get('software', {})
    
    # Add each software component
    for sw_key, sw_info in software.items():
        doc.add_paragraph(f"{sw_info.get('name', '')}:")
        doc.add_paragraph(f"  Version: {sw_info.get('version', '')}")
        if sw_info.get('comment') and sw_info.get('comment') != '_':
            doc.add_paragraph(f"  Comment: {sw_info.get('comment', '')}")
    
    # Hardware Information
    doc.add_heading("Hardware", level=2)
    hardware = config.get('system_info', {}).get('hardware', {})
    
    # Add each PC configuration
    for pc_key, pc_info in hardware.items():
        doc.add_paragraph(f"{pc_info.get('name', '')}:")
        if pc_info.get('comment_1') and pc_info.get('comment_1') != '_':
            doc.add_paragraph(f"  {pc_info.get('comment_1', '')}")
        if pc_info.get('comment_2') and pc_info.get('comment_2') != '_':
            doc.add_paragraph(f"  {pc_info.get('comment_2', '')}")

def add_environment_section(doc, config):
    """
    Add environment information section to the document.
    
    This function creates a section detailing the test environment configuration,
    including multiple parameter sets for the test environment.
    
    Args:
        doc: The document object to add content to
        config: The configuration dictionary containing environment information
    
    The section includes:
        - Multiple parameter sets (Param_1, Param_2, etc.):
            * Name
            * Type
            * Location
    """
    doc.add_heading("Test Environment", level=1)
    env = config.get('environment', {})
    
    # Add each parameter set
    for param_key, param_info in env.items():
        doc.add_heading(param_info.get('name', param_key), level=2)
        doc.add_paragraph(f"Type: {param_info.get('type', '')}")
        doc.add_paragraph(f"Location: {param_info.get('location', '')}")

def apply_document_settings(doc, config):
    """
    Apply document formatting settings from config.
    
    This function applies the document-wide formatting settings including headers,
    footers, fonts, and colors as specified in the configuration file.
    
    Args:
        doc: The document object to apply settings to
        config: The configuration dictionary containing document settings
    
    The function applies:
        - Header settings:
            * Company name (Elbit Systems)
            * Department (factory acceptance test)
            * Project name
        - Footer settings:
            * Confidentiality notice
            * Document version
        - Formatting:
            * Font families and sizes
            * Color schemes
            * Heading styles
    """
    settings = config.get('document_settings', {})
    
    # Apply header settings
    header = settings.get('header', {})
    section = doc.sections[0]
    header_paragraph = section.header.paragraphs[0]
    header_paragraph.text = f"{header.get('company_name', '')} - {header.get('department', '')} - {header.get('project', '')}"
    
    # Apply footer settings
    footer = settings.get('footer', {})
    footer_paragraph = section.footer.paragraphs[0]
    footer_paragraph.text = f"{footer.get('confidentiality', '')} | Version: {footer.get('document_version', '')}"
    
    # Apply formatting
    formatting = settings.get('formatting', {})
    fonts = formatting.get('font', {})
    colors = formatting.get('colors', {})
    
    # Set default font
    doc.styles['Normal'].font.name = fonts.get('body', 'Calibri')
    doc.styles['Normal'].font.size = Pt(fonts.get('size', {}).get('body', 11))
    
    # Set heading fonts
    for i in range(1, 3):
        style = doc.styles[f'Heading {i}']
        style.font.name = fonts.get('heading', 'Arial')
        style.font.size = Pt(fonts.get('size', {}).get(f'heading{i}', 14))

def create_doc_from_json(json_path_list, pictures=True, Type="ATP", Regular_doc_path=True):
    """
    Generate a Word document from a JSON test result file.

    This function loads test data from a JSON file and generates a Word (.docx) document summarizing the test,
    including a summary table of steps and optionally embedding images for each step.

    Parameters
    ----------
    json_path : str
        Path to the JSON file containing test data.
    pictures : bool, optional
        Whether to include images in the document (default: True).
    Type : str, optional
        Document type, e.g., "ATP" or "ATR" (default: "ATP").
    Regular_doc_path : bool, optional
        Whether to use the regular document path naming (default: True).

    Returns
    -------
    None
        The function saves the generated document to disk; it does not return a value.
    """
    try:
        # Create new document
        doc = Document()
        config = Config()
        
        # Load document configuration
        doc_config = load_doc_config()
        
        # Add title page and TOC placeholder if multiple documents
        if len(json_path_list) > 1:
            # Title page
            doc.add_heading(f"auto generate of {Type} document", level=0)
            doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_page_break()
            
            # Add system and environment information
            add_system_info_section(doc, doc_config)
            doc.add_page_break()
            add_environment_section(doc, doc_config)
            doc.add_page_break()
        
        # Apply document settings
        apply_document_settings(doc, doc_config)
        
        for json_path in json_path_list:
            # Load JSON data
            with open(json_path, "r") as f:
                data = json.load(f)

            doc.add_heading(data.get("comment1", "Test Report"), level=1)

            doc.add_heading("TestData", level=2)
            doc.add_paragraph(f"Description: {data.get('comment2', '')}")
            doc.add_paragraph(f"Starting Point: {data.get('starting_point', '')}")
            doc.add_paragraph(f"Accuracy Level: {data.get('accuracy_level', '')}")
            doc.add_paragraph(f"Timestamp: {data.get('timestamp', '')}")

            doc.add_heading("Test precondition", level=2)    
            doc.add_paragraph(f"Description: {data.get('precondition', '')}")
            
            doc.add_heading("Steps summary table", level=2)
            # Add summary table for all steps
            table = doc.add_table(rows=1, cols=5)
            table.autofit = False  # Prevent Word from auto-resizing

            column_widths = [Inches(0.5), Inches(1.7), Inches(1.7), Inches(1.0), Inches(1.0)]
            for idx, width in enumerate(column_widths):
                for cell in table.columns[idx].cells:
                    cell.width = width

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Step #'
            hdr_cells[1].text = 'Description'
            hdr_cells[2].text = 'Acceptance'
            hdr_cells[3].text = 'Result'
            hdr_cells[4].text = 'Image Name'

            for event in data["events"]:
                if event.get("action") == "Special key 'f2' pressed":
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(event.get('screenshot_counter', ''))
                    row_cells[1].text = event.get('step_desc', '')
                    row_cells[2].text = event.get('step_accep', '')
                    row_cells[3].text = event.get('step_resau', '')
                    row_cells[4].text = event.get('image_name', '')

            # After adding all rows (header and data), set width for every cell in every row
            for row in table.rows:
                for idx, width in enumerate(column_widths):
                    row.cells[idx].width = width

            def set_table_borders(table):
                """Add borders to the summary table."""
                tbl = table._tbl
                tblPr = tbl.tblPr
                borders = OxmlElement('w:tblBorders')
                for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '8')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), '000000')
                    borders.append(border)
                tblPr.append(borders)

            # After filling the table:
            set_table_borders(table)
            if pictures:
                doc.add_heading("Steps", level=2)
            
                for event in data["events"]:
                    if event.get("action") == "Special key 'f2' pressed":
                        step_title = f"{event.get('screenshot_counter', '')}: {event.get('image_name', '')}"
                        doc.add_heading(step_title, level=3)
                        doc.add_paragraph(f"Position: X= {event.get('pic_x', '')} , Y= {event.get('pic_y', '')} \t Dimension: W= {event.get('pic_width', '')} , H= {event.get('pic_height', '')}")
                        doc.add_paragraph(f"Result: {event.get('step_resau', '')}")
                        doc.add_paragraph(f"Time: {event.get('time', '')} ms")

                        # Add image if available and not "none"
                        pic_path = event.get("pic_path", "none")
                        if pic_path and pic_path.lower() != "none" and os.path.exists(pic_path):
                            try:
                                if Type == "ATR":
                                    # Calculate max width for each image
                                    section = doc.sections[0]
                                    page_width = section.page_width
                                    left_margin = section.left_margin
                                    right_margin = section.right_margin
                                    available_width = page_width - left_margin - right_margin
                                    img_width = (available_width / 2) - Inches(0.1)

                                    # Prepare diff image path
                                    base, ext = os.path.splitext(pic_path)
                                    diff_path = base + "diff_" + ext
                                    # Create a table with 2 columns for side-by-side images
                                    img_table = doc.add_table(rows=2, cols=2)
                                    img_table.autofit = False
                                    # First row: images
                                    cell1 = img_table.cell(0, 0)
                                    cell2 = img_table.cell(0, 1)
                                    run1 = cell1.paragraphs[0].add_run()
                                    run1.add_picture(pic_path, width=img_width)
                                    cell1.paragraphs[0].alignment = 1  # Center
                                    if os.path.exists(diff_path):
                                        run2 = cell2.paragraphs[0].add_run()
                                        run2.add_picture(diff_path, width=img_width)
                                        cell2.paragraphs[0].alignment = 1  # Center
                                    # Second row: captions
                                    cell1_caption = img_table.cell(1, 0)
                                    cell2_caption = img_table.cell(1, 1)
                                    cell1_caption.text = f"Figure: {event.get('image_name', '')}"
                                    cell1_caption.paragraphs[0].alignment = 1
                                    cell2_caption.text = "Figure: Diff"
                                    cell2_caption.paragraphs[0].alignment = 1
                                else:
                                    doc.add_picture(pic_path, width=Inches(4))
                                    caption = event.get('image_name', '')
                                    if caption:
                                        last_paragraph = doc.paragraphs[-1]
                                        last_paragraph.alignment = 1  # Center the image
                                        doc.add_paragraph(f"Figure: {caption}").alignment = 1  # Center the caption
                            except Exception as e:
                                doc.add_paragraph(f"Could not add image: {pic_path} ({e})")

            def add_page_number(paragraph):
                """Add a page number field to the given paragraph."""
                run = paragraph.add_run()
                fldChar1 = OxmlElement('w:fldChar')
                fldChar1.set(qn('w:fldCharType'), 'begin')
                instrText = OxmlElement('w:instrText')
                instrText.text = 'PAGE'
                fldChar2 = OxmlElement('w:fldChar')
                fldChar2.set(qn('w:fldCharType'), 'end')
                run._r.append(fldChar1)
                run._r.append(instrText)
                run._r.append(fldChar2)

            section = doc.sections[0]
             
            # # Header
            # header = section.header
            # header_paragraph = header.paragraphs[0]
            # header_paragraph.text = data.get(f"auto generate of {Type} document", "Test Report")

            # # Footer
            # footer = section.footer
            # footer_paragraph = footer.paragraphs[0]
            # footer_paragraph.text = "Page "
            # add_page_number(footer_paragraph)
            # footer_paragraph.add_run(f" | Date: {datetime.now().strftime('%Y-%m-%d')}")

            doc.add_page_break()

            


        if Regular_doc_path == True:
            base = os.path.splitext(os.path.basename(json_path))[0]
            doc_path = os.path.join(os.path.dirname(json_path), f"{base}.docx")
            doc.save(doc_path)
            print(f"Word document created: {doc_path}")
            return True
        else:
            # Get paths from config
            paths_config = config.get('paths', {})
            db_path = paths_config.get('db_path', os.path.join(project_root, "DB"))
            if Type == "ATP":
                doc_path = os.path.join(db_path, f"ATP_document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
            else:
                doc_path = os.path.join(db_path, f"ATR_document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
            doc.save(doc_path)
            print(f"Word document created: {doc_path}")
            
            # Automatically open the document
            try:
                if sys.platform == 'win32':
                    os.startfile(doc_path)
                elif sys.platform == 'darwin':  # macOS
                    subprocess.run(['open', doc_path])
                else:  # Linux
                    subprocess.run(['xdg-open', doc_path])
                print("Document opened successfully")
            except Exception as e:
                print(f"Could not open document automatically: {e}")
            
            return True
        return False

    except Exception as e:
        print(f"Error creating document: {e}")
        return False

if __name__ == "__main__":
    # Example usage
    json_path = r"DB\Result\20250520_220950_map scroll slow\Result_map scroll slow.json"
    create_doc_from_json(json_path, pictures=True, Type="ATR")
